from docx import Document

from docxtpl import DocxTemplate
import sys

inform = sys.stdin.read().split("\n")
document = Document()
for i in range(len(inform))
    document.add_heading("Приглошение", 0)
    p = document.add_paragraph(f"Прив {inform[i]}")
    p = document.add_paragraph(f"приглашаю в {inform[1]}{inform[0]}")
    document.add_page_break()
document.save("res.docx")
