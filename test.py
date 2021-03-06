from PIL import Image, ImageDraw


def picture(file_name, width, height, sky_color="#87CEEB", ocean_color="#017B92", boat_color="#874535",
            sail_color="#FFFFFF", sun_color="#FFCF40"):
    im = Image.new("RGB", (width, height))
    drawer = ImageDraw.Draw(im)

    drawer.rectangle(((0, 0), (width, int(height * 0.8))), sky_color)
    drawer.rectangle(((0, int(height * 0.8)), (width, height)),
                     ocean_color)
    drawer.ellipse((
        (int(0.8 * width), -int(0.2 * height)),
        (int(1.2 * width), int(0.2 * height))),
        sun_color)
    drawer.polygon(((int(0.25 * width), int(0.65 * height)),
                    (int(0.49 * width), int(0.65 * height)),
                    (int(0.49 * width), int(0.3 * height)),
                    (int(0.51 * width), int(0.3 * height)),
                    (int(0.51 * width), int(0.65 * height)),
                    (int(0.75 * width), int(0.65 * height)),
                    (int(0.7 * width), int(0.85 * height)),
                    (int(0.3 * width), int(0.85 * height))), boat_color)
    drawer.polygon(((int(0.51 * width + 1), int(0.3 * height)),
                    (int(0.66 * width + 1), int(0.45 * height)),
                    int(0.51 * width + 1), int(0.6 * height)), sail_color)
    im.save(file_name)


picture('rss.jpg', 1000, 800)

from docx import Document


def markdown_to_docx(text):
    document = Document()
    lines = text.split('\n')
    document.add_heading(lines[0], 0)
    for line in lines[1:]:
        if line:
            if line[:7].count('#') == 1:
                document.add_heading(line[2:], level=1)
            elif line[:7].count('#') == 2:
                document.add_heading(line[3:], level=2)
            elif line[:7].count('#') == 3:
                document.add_heading(line[4:], level=3)
            elif line[:7].count('#') == 4:
                document.add_heading(line[5:], level=4)
            elif line[:7].count('#') == 5:
                document.add_heading(line[6:], level=5)
            elif line[:7].count('#') == 6:
                document.add_heading(line[7:], level=6)
            elif str(line[:2]) == '- ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '* ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '+ ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and line[1] == '.':
                document.add_paragraph(line[3:], style='List Number')
            elif line[:3].count('_') == 1 or line[:3].count('*') == 1:
                document.add_paragraph().add_run(line[1:-1]).italic = True
            elif line[:3].count('_') == 2 or line[:3].count('*') == 2:
                document.add_paragraph().add_run(line[2:-2]).bold = True
            elif line[:3].count('_') == 3 or line[:3].count('*') == 3:
                runner = document.add_paragraph().add_run(line[3:-3])
                runner.bold = True
                runner.italic = True
            else:
                document.add_paragraph(line)
        else:
            document.add_paragraph()
    document.save('res.docx')








